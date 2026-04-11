#!/usr/bin/env python3
"""Convert OpenClaw Markdown document to Word with embedded images."""

import re
import os
import shutil
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

# Paths
md_path = "/Users/Zhuanz/Desktop/yp-nx-dashboard/docs/OpenClaw-Deployment-Guide.md"
assets_src = "/Users/Zhuanz/Desktop/yp-nx-dashboard/docs/assets"
export_dir = "/Users/Zhuanz/Desktop/yp-nx-dashboard/docs/OpenClaw_Export"
docx_path = os.path.join(export_dir, "OpenClaw-Doctor-部署指南.docx")

# Create export directory
os.makedirs(export_dir, exist_ok=True)

# Copy assets to export folder
assets_dst = os.path.join(export_dir, "assets")
if os.path.exists(assets_src):
    if os.path.exists(assets_dst):
        shutil.rmtree(assets_dst)
    shutil.copytree(assets_src, assets_dst)

# Read markdown
with open(md_path, 'r', encoding='utf-8') as f:
    content = f.read()

# Create document
doc = Document()

# Set default font for Chinese
style = doc.styles['Normal']
style.font.name = 'Microsoft YaHei'
style._element.rPr.rFonts.set('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}eastAsia', 'Microsoft YaHei')

def add_heading(doc, text, level):
    """Add heading with proper styling."""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}eastAsia', 'Microsoft YaHei')
    return heading

def add_paragraph(doc, text, bold=False, italic=False):
    """Add paragraph with styling."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}eastAsia', 'Microsoft YaHei')
    return p

def add_code_block(doc, code, language=''):
    """Add code block."""
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(10)
    p.paragraph_format.left_indent = Inches(0.3)
    return p

def add_image(doc, image_path):
    """Add image from relative path."""
    full_path = os.path.join(export_dir, image_path)
    if os.path.exists(full_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(full_path, width=Inches(5.5))
        return p
    return None

def add_table(doc, headers, rows):
    """Add table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'

    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            for run in p.runs:
                run.bold = True

    # Data rows
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, cell in enumerate(row):
            cells[ci].text = str(cell)

    return table

# Parse and convert
lines = content.split('\n')
i = 0
in_code_block = False
code_block_content = []
table_headers = []
table_rows = []
in_table = False

while i < len(lines):
    line = lines[i]

    # Code blocks
    if line.strip().startswith('```'):
        if not in_code_block:
            in_code_block = True
            code_block_content = []
        else:
            # End code block
            add_code_block(doc, '\n'.join(code_block_content))
            in_code_block = False
        i += 1
        continue

    if in_code_block:
        code_block_content.append(line)
        i += 1
        continue

    # Tables
    if '|' in line and line.strip().startswith('|'):
        # Parse table row
        cells = [c.strip() for c in line.split('|')[1:-1]]
        if not any('---' in c for c in cells):
            if in_table and not table_headers:
                table_headers = cells
            elif in_table:
                table_rows.append(cells)
            else:
                table_headers = cells
                in_table = True
        i += 1
        continue
    else:
        if in_table and table_headers:
            add_table(doc, table_headers, table_rows)
            table_headers = []
            table_rows = []
        in_table = False

    # Images
    img_match = re.search(r'!\[([^\]]*)\]\(([^)]+)\)', line)
    if img_match:
        alt_text, img_path = img_match.groups()
        # Use relative path from export dir
        if not img_path.startswith('assets/'):
            img_path = 'assets/' + os.path.basename(img_path)
        add_image(doc, img_path)
        i += 1
        continue

    # Headings
    if line.startswith('# '):
        add_heading(doc, line[2:], 0)
    elif line.startswith('## '):
        add_heading(doc, line[3:], 1)
    elif line.startswith('### '):
        add_heading(doc, line[4:], 2)
    elif line.startswith('#### '):
        add_heading(doc, line[5:], 3)
    # Horizontal rule
    elif line.strip() == '---':
        doc.add_paragraph('─' * 50)
    # Empty line
    elif line.strip() == '':
        pass
    # List items
    elif line.strip().startswith('- ') or line.strip().startswith('* '):
        add_paragraph(doc, line.strip()[2:])
    # Bold/Italic
    elif '**' in line:
        parts = re.split(r'(\*\*[^*]+\*\*)', line)
        p = doc.add_paragraph()
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.bold = True
            else:
                run = p.add_run(part)
            run.font.name = 'Microsoft YaHei'
    # Regular paragraph
    else:
        add_paragraph(doc, line)

    i += 1

# Save
doc.save(docx_path)
print(f"Document saved to: {docx_path}")
print(f"Assets folder: {assets_dst}")
